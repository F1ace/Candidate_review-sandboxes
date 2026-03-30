from __future__ import annotations

import hashlib
import re
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any
from uuid import uuid4

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader
from sqlalchemy.orm import Session

from .. import models
from ..config import settings
from .object_storage import storage_service


TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".py",
    ".sql",
    ".json",
    ".csv",
    ".yml",
    ".yaml",
    ".log",
}
HTML_EXTENSIONS = {".html", ".htm"}
DOCX_EXTENSIONS = {".docx"}
PDF_EXTENSIONS = {".pdf"}


def _decode_text_bytes(payload: bytes) -> str:
    encodings = ("utf-8-sig", "utf-8", "cp1251", "latin-1")
    for encoding in encodings:
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    return payload.decode("utf-8", errors="ignore")


def extract_text_from_bytes(filename: str, payload: bytes, content_type: str | None = None) -> str:
    suffix = Path(filename).suffix.lower()

    if suffix in TEXT_EXTENSIONS:
        return _decode_text_bytes(payload).strip()

    if suffix in HTML_EXTENSIONS or (content_type or "").lower().startswith("text/html"):
        html = _decode_text_bytes(payload)
        return BeautifulSoup(html, "html.parser").get_text("\n", strip=True).strip()

    if suffix in PDF_EXTENSIONS or (content_type or "").lower() == "application/pdf":
        reader = PdfReader(BytesIO(payload))
        parts = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(part.strip() for part in parts if part.strip()).strip()

    if suffix in DOCX_EXTENSIONS or (
        content_type or ""
    ).lower() == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        document = DocxDocument(BytesIO(payload))
        return "\n".join(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()).strip()

    raise ValueError(f"Unsupported document type: {filename}")


def chunk_text(text: str, *, chunk_size: int | None = None, overlap: int | None = None) -> list[dict[str, Any]]:
    normalized = (text or "").strip()
    if not normalized:
        return []

    target_size = int(chunk_size or settings.rag_chunk_size)
    target_overlap = int(overlap or settings.rag_chunk_overlap)
    if target_size <= 0:
        raise ValueError("chunk_size must be positive")
    if target_overlap < 0 or target_overlap >= target_size:
        raise ValueError("chunk_overlap must be within [0, chunk_size)")

    words = list(re.finditer(r"\S+", normalized))
    if not words:
        return []

    chunks: list[dict[str, Any]] = []
    start_idx = 0
    chunk_index = 0
    step = max(1, target_size - target_overlap)

    while start_idx < len(words):
        end_idx = min(len(words), start_idx + target_size)
        char_start = words[start_idx].start()
        char_end = words[end_idx - 1].end()
        chunk_content = normalized[char_start:char_end].strip()
        if chunk_content:
            chunks.append(
                {
                    "chunk_index": chunk_index,
                    "content": chunk_content,
                    "content_length": len(chunk_content),
                    "token_count": end_idx - start_idx,
                    "char_start": char_start,
                    "char_end": char_end,
                }
            )
            chunk_index += 1
        if end_idx >= len(words):
            break
        start_idx += step

    return chunks


def _build_object_key(corpus_id: int, checksum_sha256: str, filename: str) -> str:
    safe_name = Path(filename).name or "document.txt"
    return f"corpora/{corpus_id}/{checksum_sha256[:12]}-{uuid4().hex}-{safe_name}"


def ingest_document_bytes(
    *,
    db: Session,
    corpus_id: int,
    filename: str,
    payload: bytes,
    content_type: str,
    metadata: dict[str, Any] | None = None,
    extracted_text: str | None = None,
) -> models.Document:
    checksum_sha256 = hashlib.sha256(payload).hexdigest()
    object_key = _build_object_key(corpus_id, checksum_sha256, filename)
    storage_info = storage_service.put_bytes(object_key=object_key, payload=payload, content_type=content_type)

    text = (extracted_text or extract_text_from_bytes(filename, payload, content_type)).strip()
    if not text:
        raise ValueError("Extracted text is empty")

    document = models.Document(
        rag_corpus_id=corpus_id,
        filename=filename,
        content=text,
        content_type=content_type,
        storage_bucket=str(storage_info["bucket"]),
        object_key=str(storage_info["object_key"]),
        size_bytes=int(storage_info["size_bytes"]),
        checksum_sha256=checksum_sha256,
        status="ready",
        ingested_at=datetime.utcnow(),
        meta=metadata or {},
    )
    db.add(document)
    db.flush()

    chunks = chunk_text(text)
    for chunk in chunks:
        db.add(
            models.DocumentChunk(
                document_id=document.id,
                chunk_index=chunk["chunk_index"],
                content=chunk["content"],
                content_length=chunk["content_length"],
                token_count=chunk["token_count"],
                char_start=chunk["char_start"],
                char_end=chunk["char_end"],
                meta={"source_filename": filename},
            )
        )

    db.commit()
    db.refresh(document)
    return document
